"""
Conversation -> Word document export.

Builds a .docx that is faithful to what the chat displayed: each question and
answer in order, charts embedded as PNG images (captured client-side from the
exact canvases the user saw), and an optional appendix with the query and
calculation traces. Styled with the GA navy palette and the GA logo.

Pure python-docx — no LLM calls, no changes to the analytical core.
"""

import base64
import io
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x0A, 0x16, 0x28)
NAVY_MID = RGBColor(0x1B, 0x3A, 0x6B)
GREY = RGBColor(0x5F, 0x5E, 0x5A)
SILVER = RGBColor(0x8A, 0x9B, 0xB0)

LOGO_PATH = os.path.join("static", "images", "ga_logo.png")


def _shade(cell, hex_color: str):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def _decode_png(data_url: str) -> Optional[bytes]:
    """Accept a canvas dataURL ('data:image/png;base64,....') or raw base64."""
    try:
        if "," in data_url:
            data_url = data_url.split(",", 1)[1]
        return base64.b64decode(data_url)
    except Exception:
        return None


def build_conversation_docx(payload: dict) -> bytes:
    """
    payload = {
      "mode": "standard" | "variance",
      "filename": str,               # uploaded data file name(s)
      "meta": str,                   # e.g. "12,480 rows x 14 columns"
      "user": str | None,
      "include_debug": bool,
      "messages": [ { "question": str, "answer": str,
                      "chart_png": dataURL | None,
                      "chart_title": str | None,
                      "debug": [ {"label": str, "content": str}, ... ] } ]
    }
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Header: logo + title ──────────────────────────────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = True
    logo_cell, title_cell = header_tbl.rows[0].cells
    if os.path.exists(LOGO_PATH):
        run = logo_cell.paragraphs[0].add_run()
        try:
            run.add_picture(LOGO_PATH, width=Inches(0.55))
        except Exception:
            pass
    p = title_cell.paragraphs[0]
    r = p.add_run("Chat with Data — conversation report")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY

    mode_label = "Variance analysis" if payload.get("mode") == "variance" else "Standard analysis"
    meta_lines = [
        f"{mode_label}  ·  {payload.get('filename') or 'uploaded data'}",
        (payload.get("meta") or "").strip(),
        f"Exported {datetime.now().strftime('%B %d, %Y %H:%M')}"
        + (f"  ·  {payload['user']}" if payload.get("user") else ""),
    ]
    for line in [l for l in meta_lines if l]:
        mp = doc.add_paragraph()
        mr = mp.add_run(line)
        mr.font.size = Pt(9)
        mr.font.color.rgb = GREY
        mp.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ── Q&A blocks ────────────────────────────────────────────────────────────
    messages = payload.get("messages") or []
    for i, m in enumerate(messages, start=1):
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(10)
        qr = qp.add_run(f"Q{i} · {m.get('question', '').strip()}")
        qr.font.size = Pt(11)
        qr.font.bold = True
        qr.font.color.rgb = NAVY_MID

        answer = (m.get("answer") or "").strip()
        if answer:
            for para_text in [t for t in answer.split("\n") if t.strip()]:
                ap = doc.add_paragraph()
                ar = ap.add_run(para_text.strip())
                ar.font.size = Pt(10.5)
                ap.paragraph_format.space_after = Pt(4)

        png = _decode_png(m.get("chart_png") or "") if m.get("chart_png") else None
        if png:
            if m.get("chart_title"):
                cp = doc.add_paragraph()
                cr = cp.add_run(m["chart_title"])
                cr.font.size = Pt(9)
                cr.font.color.rgb = GREY
                cp.paragraph_format.space_after = Pt(2)
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(5.6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    # ── Appendix: query & calculation traces ──────────────────────────────────
    if payload.get("include_debug"):
        doc.add_page_break()
        hp = doc.add_paragraph()
        hr = hp.add_run("Appendix A — query and calculation traces")
        hr.font.size = Pt(13)
        hr.font.bold = True
        hr.font.color.rgb = NAVY

        for i, m in enumerate(messages, start=1):
            debug = m.get("debug") or []
            if not debug:
                continue
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(8)
            qr = qp.add_run(f"Q{i} · {m.get('question', '').strip()}")
            qr.font.size = Pt(10)
            qr.font.bold = True
            qr.font.color.rgb = NAVY_MID
            for step in debug:
                lp = doc.add_paragraph()
                lr = lp.add_run(str(step.get("label", "")).upper())
                lr.font.size = Pt(8)
                lr.font.bold = True
                lr.font.color.rgb = SILVER
                lp.paragraph_format.space_after = Pt(1)

                tbl = doc.add_table(rows=1, cols=1)
                cell = tbl.rows[0].cells[0]
                _shade(cell, "F5F6F8")
                cp = cell.paragraphs[0]
                content = str(step.get("content", ""))
                if len(content) > 6000:
                    content = content[:6000] + "\n... [truncated]"
                cr = cp.add_run(content)
                cr.font.name = "Consolas"
                cr.font.size = Pt(8)
                doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Footer with page numbers ──────────────────────────────────────────────
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Chat with Data — conversation report · page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = SILVER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer_p._p.append(fld)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
